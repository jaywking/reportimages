"""Word document generation utilities."""
from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from datetime import datetime

from image_utils import ImageEntry, resize_image_to_width


IMAGE_WIDTH_INCHES_TO_PIXELS = {
    2.0: 600,
    2.25: 750,
    2.5: 900,
}


def _ensure_setting(settings_element, tag: str, val: str) -> None:
    existing = settings_element.find(qn(tag))
    if existing is None:
        element = OxmlElement(tag)
        element.set(qn("w:val"), val)
        settings_element.append(element)
    else:
        existing.set(qn("w:val"), val)


def apply_image_preferences(document: Document) -> None:
    """Disable image compression and set high-fidelity DPI."""

    settings = document.settings.element
    _ensure_setting(settings, "w:doNotCompressPictures", "true")
    # High-fidelity resolution is represented by 0 per WordprocessingML schema.
    _ensure_setting(settings, "w:defaultImageDpi", "0")
    # Allow single-click to follow hyperlinks (no Ctrl modifier required).
    _ensure_setting(settings, "w:doNotUseCtrlClick", "true")


def _add_hyperlinked_picture(document: Document, image_stream, width_inches: float, hyperlink: str):
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(image_stream, width=Inches(width_inches))

    rel_id = document.part.relate_to(hyperlink, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink_el = OxmlElement("w:hyperlink")
    hyperlink_el.set(qn("r:id"), rel_id)
    hyperlink_el.set(qn("w:history"), "1")

    run_element = run._element
    run_element.getparent().remove(run_element)
    hyperlink_el.append(run_element)
    paragraph._p.append(hyperlink_el)

    # Also apply the hyperlink directly to the picture non-visual properties to
    # help Word recognize the image as a clickable link.
    for el in run_element.iter():
        tag = el.tag
        if isinstance(tag, str) and tag.endswith("}cNvPr"):
            hlink_click = OxmlElement("a:hlinkClick")
            hlink_click.set(qn("r:id"), rel_id)
            el.append(hlink_click)
            break


def build_document(entries: Iterable[ImageEntry], width_inches: float, output_path: Path) -> None:
    """Create a Word document containing the provided images.

    Images are resized in-memory before insertion. One image is placed per
    paragraph, wrapped in a hyperlink pointing to the original asset.
    """

    target_width_px = IMAGE_WIDTH_INCHES_TO_PIXELS.get(width_inches)
    if target_width_px is None:
        raise ValueError(f"Unsupported image width: {width_inches}")

    document = Document()
    # Use a naive UTC timestamp for compatibility with python-docx expectations.
    now = datetime.utcnow()
    core_props = document.core_properties
    core_props.created = now
    core_props.modified = now

    apply_image_preferences(document)

    for entry in entries:
        resized_stream = resize_image_to_width(entry.path, target_width_px)
        _add_hyperlinked_picture(document, resized_stream, width_inches, entry.hyperlink)

    document.save(output_path)
